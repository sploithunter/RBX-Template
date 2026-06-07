#!/usr/bin/env python3
"""Build a printable Word doc for PET_REALM_ICONS_AND_POWERS.md with embedded blue icons."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs/PET_REALM_ICONS_AND_POWERS.md"
ICON_DIR = ROOT / "assets/ui/blue_icons"
GEM_SHEET = ROOT / "assets/ui/gems_all_colors.png"
OUT_PATH = ROOT / "docs/PET_REALM_ICONS_AND_POWERS.docx"

ICON_THUMB = Inches(0.45)
SHEET_WIDTH = Inches(6.5)


def symbol_from_cell(cell: str) -> str | None:
    match = re.search(r"`([a-z0-9_]+)`", cell)
    return match.group(1) if match else None


def add_icon_cell(cell, symbol: str | None) -> None:
    if not symbol:
        cell.text = ""
        return
    path = ICON_DIR / f"{symbol}.png"
    if not path.is_file():
        cell.text = f"(no art: {symbol})"
        return
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=ICON_THUMB)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if not line.startswith("|"):
        return []
    parts = [part.strip() for part in line.strip("|").split("|")]
    return parts


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells if c)


def add_markdown_table(
    doc: Document,
    rows: list[list[str]],
    *,
    icon_column: int | None = None,
) -> None:
    if len(rows) < 2:
        return
    header = rows[0]
    body = [row for row in rows[1:] if not is_separator_row(row)]
    if icon_column is not None:
        header = ["Icon", *header]
        body = [["", *row] for row in body]

    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    for col, text in enumerate(header):
        table.rows[0].cells[col].text = text.replace("`", "")

    for row_idx, row in enumerate(body, start=1):
        for col, text in enumerate(row):
            clean = text.replace("`", "")
            if icon_column is not None and col == 0:
                symbol = symbol_from_cell(row[icon_column])
                add_icon_cell(table.rows[row_idx].cells[col], symbol)
            else:
                table.rows[row_idx].cells[col].text = clean
    doc.add_paragraph()


def add_heading(doc: Document, line: str) -> None:
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)


def add_paragraph(doc: Document, line: str) -> None:
    text = line.strip()
    if not text:
        return
    if text == "---":
        doc.add_paragraph()
        return
    if text.startswith("> "):
        p = doc.add_paragraph(text[2:].strip())
        p.style = "Intense Quote"
        return
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(9)
        else:
            p.add_run(part)


def convert_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    table_buf: list[list[str]] = []
    symbol_catalog = False
    misleading = False

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("|"):
            cells = parse_table_row(line)
            if cells:
                table_buf.append(cells)
            i += 1
            continue

        if table_buf:
            icon_col = None
            if symbol_catalog and table_buf[0] and table_buf[0][0].lower().startswith("symbol"):
                icon_col = 0
            elif misleading and table_buf[0] and "asset key" in table_buf[0][0].lower():
                icon_col = 0
            add_markdown_table(doc, table_buf, icon_column=icon_col)
            table_buf = []
            symbol_catalog = False
            misleading = False

        if line.startswith("#"):
            add_heading(doc, line)
            title = line.lstrip("# ").strip()
            if title.startswith("A.3"):
                symbol_catalog = True
            if title.startswith("A.4"):
                misleading = True
            i += 1
            continue

        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        add_paragraph(doc, line)
        i += 1

    if table_buf:
        add_markdown_table(doc, table_buf)


def add_contact_sheets(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix — Icon contact sheets", level=1)

    blue_sheet = ICON_DIR / "contact_sheet.png"
    if blue_sheet.is_file():
        doc.add_heading("Power symbols (blue set)", level=2)
        doc.add_picture(str(blue_sheet), width=SHEET_WIDTH)

    if GEM_SHEET.is_file():
        doc.add_heading("Gem meshes (all colors)", level=2)
        doc.add_picture(str(GEM_SHEET), width=SHEET_WIDTH)


def main() -> int:
    if not MD_PATH.is_file():
        print(f"Missing: {MD_PATH}", file=sys.stderr)
        return 1

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(10)

    doc.add_heading("Pet Realm — Icons, Colors & Powers", level=0)
    doc.add_paragraph(
        "Printable reference with embedded blue icon art. "
        "Generated from docs/PET_REALM_ICONS_AND_POWERS.md."
    )

    convert_markdown(doc, MD_PATH.read_text(encoding="utf-8"))
    add_contact_sheets(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
